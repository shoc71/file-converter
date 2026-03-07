# imports 
import os
from PIL import Image
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import markdown
from weasyprint import HTML
import fitz
from pypdf import PdfReader
from markdownify import markdownify as md 
from pathlib import Path


def doc_to_txt(input_path, output_name):
    doc = Document(input_path)

    text = "\n".join(p.text for p in doc.paragraphs)

    with open(output_name, "w") as f:
        f.write(text)

def png_to_jpg(input_path, output_name):
    img = Image.open(input_path).convert("RGB")
    img.save(output_name, quality=100)

def txt_to_pdf(input_path, output_name):
    doc = SimpleDocTemplate(output_name, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    with open(input_path, encoding="utf-8") as f:
        for line in f:
            elements.append(Paragraph(line.strip(), styles["Normal"]))
            elements.append(Spacer(1, 0.2 * inch))

    doc.build(elements)

def md_to_pdf(input_path, output_name):
    with open(input_path) as f:
        md_text = f.read()
    
    html_body = markdown.markdown(
        md_text,
        extensions=["extra", "codehilite", "tables", "toc"]
        )

    full_html = f"""
        <html>
        <head>
        <meta ="utf-8">
        <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            margin: 40px;
        }}
        h1, h2, h3 {{
            color: #333;
        }}
        code {{
            background: #f4f4f4;
            padding: 2px 4px;
        }}
        pre {{
            background: #f4f4f4;
            padding: 10px;
            overflow-x: auto;
        }}
        </style>
        </head>
        <body>
        {html_body}
        </body>
        </html>
    """

    HTML(string=full_html).write_pdf(output_name)

def html_to_pdf(input_path, output_name):
    HTML(input_path).write_pdf(output_name)

def png_to_pdf(input_path, output_name):
    images = [input_path]

    pil_images = [Image.open(img).convert("RGB") for img in images]

    pil_images[0].save(
        output_name,
        save_all=True
    )

def pdf_to_html(input_path, output_name):
    doc = fitz.open(input_path)
    html_parts = []

    for page in doc:
        html_parts.append(page.get_text("xhtml"))

    full_html = f"""<!DOCTYPE html>
        <html>
        <head>
            <meta ="utf-8">
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    margin: 40px;
                }}
                img {{
                    max-width: 100%;
                }}
            </style>
        </head>
        <body>
        {''.join(html_parts)}
        </body>
        </html>
    """

    Path(output_name).write_text(full_html, encoding="utf-8")

def pdf_to_txt(input_path, output_name):
    reader = PdfReader(input_path)

    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    
    with open(output_name, "w", encoding="utf-8") as f:
        f.write(text)

def pdf_to_png(input_path, output_name):
    doc = fitz.open(input_path)

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap(dpi=300)
        pix.save(f"./output/page_{page_num+1}.png")

def txt_to_docx(input_path, output_name):
    doc = Document()

    with open(input_path, encoding="utf-8") as f:
        for line in f:
            doc.add_paragraph(line.rstrip())

    doc.save(output_name)

# technically pdf -> html -> md
def pdf_to_md(input_pdf, output_md, image_dir="images"):
    doc = fitz.open(input_pdf)
    Path(image_dir).mkdir(exist_ok=True)

    html_parts = []

    for page_index, page in enumerate(doc):
        # get layout HTML (includes positioning & formatting)
        html_parts.append(page.get_text("html"))

        # extract images
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            pix = fitz.Pixmap(doc, xref)

            img_name = f"img_{page_index+1}_{img_index+1}.png"
            img_path = Path(image_dir) / img_name

            if pix.n < 5:
                pix.save(img_path)
            else:
                rgb_pix = fitz.Pixmap(fitz.csRGB, pix)
                rgb_pix.save(img_path)
                rgb_pix = None

            pix = None  # free memory

            # add markdown image reference
            html_parts.append(f'<p><img src="{image_dir}/{img_name}"></p>')

    # convert combined HTML to markdown
    markdown_text = md("".join(html_parts))

    Path(output_md).write_text(markdown_text, encoding="utf-8")

def main():
    print("hello world")

def test():
    print("running tests now")

    OUTPUT_PATH = "./output"
    if not os.path.exists(OUTPUT_PATH):
        os.mkdir(OUTPUT_PATH)

    png_to_jpg("input/sonic.png", "output/output-sonic.jpg")
    doc_to_txt("input/cat.docx", "output/output-cat.txt") # case specific
    txt_to_pdf("input/hi.txt", "output/output-hi.pdf")
    md_to_pdf("input/hihi.md", "output/output-hihi.pdf")
    html_to_pdf("input/example.html", "output/output-example.pdf")
    png_to_pdf("input/sonic.png", "output/output-sonic.pdf")
    pdf_to_html("input/online.pdf", "output/output-online.html")
    pdf_to_txt("input/online.pdf", "output/output-online.txt")
    pdf_to_png("input/online.pdf", "output/output-online.png")
    txt_to_docx("input/hi.txt", "output/output-hi.docx")
    pdf_to_md("input/online.pdf", "output/output-online.md")

    print("test complete")

if __name__ == "__main__":
    main()
    test()